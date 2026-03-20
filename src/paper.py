from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START

OUTPUT = r"d:\GRA\EDSR-PyTorch-master\thesis_template_pfdn.docx"

def _enable_update_fields_on_open(doc):
    settings = doc.settings.element
    # 先删除已有的同名设置，避免重复
    for node in settings.findall(qn("w:updateFields")):
        settings.remove(node)
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def set_run_font(run, cn="宋体", en="Times New Roman", size=Pt(12), bold=False):
    run.font.name = en
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn)
    run.font.size = size
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def _safe_get_style(doc, names):
    for name in names:
        try:
            return doc.styles[name]
        except KeyError:
            continue
    return doc.styles["Normal"]


def _set_style_font(doc, style_names, cn, en, size_pt, bold=False):
    style = _safe_get_style(doc, style_names)
    style.font.name = en
    style._element.rPr.rFonts.set(qn("w:eastAsia"), cn)
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)


def _setup_section_page(sec):
    sec.top_margin = Mm(25)
    sec.bottom_margin = Mm(20)
    sec.left_margin = Mm(25)
    sec.right_margin = Mm(20)
    sec.header_distance = Mm(15)
    sec.footer_distance = Mm(12)


def _set_section_page_number_format(section, fmt="decimal", start=1):
    sectPr = section._sectPr
    for child in list(sectPr):
        if child.tag == qn("w:pgNumType"):
            sectPr.remove(child)
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:fmt"), fmt)  # upperRoman / decimal
    pg_num_type.set(qn("w:start"), str(start))
    sectPr.append(pg_num_type)


def _add_page_number_to_footer(doc, section):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()

    # 直接用传入的 doc，避免依赖 section 私有属性
    p.style = _safe_get_style(doc, ["Footer", "页脚"])

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Pt(0)
    p.text = ""

    # PAGE 域（复杂域，兼容性高）
    r_begin = p.add_run()
    set_run_font(r_begin, cn="宋体", en="Times New Roman", size=Pt(10.5), bold=False)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r_begin._r.append(fld_begin)

    r_instr = p.add_run()
    set_run_font(r_instr, cn="宋体", en="Times New Roman", size=Pt(10.5), bold=False)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE \\* MERGEFORMAT "
    r_instr._r.append(instr)

    r_sep = p.add_run()
    set_run_font(r_sep, cn="宋体", en="Times New Roman", size=Pt(10.5), bold=False)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r_sep._r.append(fld_sep)

    r_text = p.add_run("1")
    set_run_font(r_text, cn="宋体", en="Times New Roman", size=Pt(10.5), bold=False)

    r_end = p.add_run()
    set_run_font(r_end, cn="宋体", en="Times New Roman", size=Pt(10.5), bold=False)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_end._r.append(fld_end)


def _insert_toc_field(doc, max_level=2):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Pt(0)

    r_begin = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r_begin._r.append(fld_begin)

    r_instr = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f'TOC \\o "1-{max_level}" \\h \\z \\u'
    r_instr._r.append(instr)

    r_sep = p.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r_sep._r.append(fld_sep)

    r_hint = p.add_run("（右键目录选择“更新域”）")
    set_run_font(r_hint, size=Pt(12))

    r_end = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_end._r.append(fld_end)


def add_paragraph(doc, text="", align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, first_line_indent_chars=2):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Pt(first_line_indent_chars * 12) if first_line_indent_chars else Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=Pt(12), bold=False)
    return p


def add_title(doc, text, size_pt=18):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(text)
    set_run_font(r, cn="黑体", en="Times New Roman", size=Pt(size_pt), bold=False)
    return p


def _add_heading(doc, text, level, align=WD_PARAGRAPH_ALIGNMENT.LEFT):
    style = _safe_get_style(doc, [f"Heading {level}", f"标题 {level}"])
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Pt(0)
    # 不再覆盖 run 字体，完全采用样式，避免冲突
    p.add_run(text)
    return p


def add_heading_lv1(doc, text):
    return _add_heading(doc, text, level=1, align=WD_PARAGRAPH_ALIGNMENT.CENTER)


def add_heading_lv2(doc, text):
    return _add_heading(doc, text, level=2, align=WD_PARAGRAPH_ALIGNMENT.LEFT)


def add_heading_lv3(doc, text):
    return _add_heading(doc, text, level=3, align=WD_PARAGRAPH_ALIGNMENT.LEFT)


def add_heading_lv4(doc, text):
    return _add_heading(doc, text, level=4, align=WD_PARAGRAPH_ALIGNMENT.LEFT)


def add_keywords(doc, cn=True):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Pt(0)

    if cn:
        r1 = p.add_run("关键词：")
        set_run_font(r1, cn="黑体", size=Pt(12), bold=False)
        r2 = p.add_run("图像超分辨率；轻量化网络；信息蒸馏；注意力机制；特征融合")
        set_run_font(r2, size=Pt(12), bold=False)
    else:
        r1 = p.add_run("Keywords: ")
        set_run_font(r1, cn="Times New Roman", en="Times New Roman", size=Pt(12), bold=False)
        r2 = p.add_run("image super-resolution; lightweight network; information distillation; attention mechanism; feature fusion")
        set_run_font(r2, cn="Times New Roman", en="Times New Roman", size=Pt(12), bold=False)


def add_chapter_page_break(doc):
    doc.add_page_break()


def configure_styles(doc):
    # 正文
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.font.bold = False
    normal.font.color.rgb = RGBColor(0, 0, 0)

    # 标题（用于大纲/目录）
    _set_style_font(doc, ["Heading 1", "标题 1"], cn="黑体", en="Times New Roman", size_pt=18, bold=False)
    # 节标题：小三宋体加粗
    _set_style_font(doc, ["Heading 2", "标题 2"], cn="宋体", en="Times New Roman", size_pt=15, bold=True)
    _set_style_font(doc, ["Heading 3", "标题 3"], cn="黑体", en="Times New Roman", size_pt=14, bold=False)
    _set_style_font(doc, ["Heading 4", "标题 4"], cn="黑体", en="Times New Roman", size_pt=12, bold=False)

    # 目录样式：TOC1黑体，其余宋体
    _set_style_font(doc, ["TOC 1", "目录 1"], cn="黑体", en="Times New Roman", size_pt=12, bold=False)
    _set_style_font(doc, ["TOC 2", "目录 2"], cn="宋体", en="Times New Roman", size_pt=12, bold=False)
    _set_style_font(doc, ["TOC 3", "目录 3"], cn="宋体", en="Times New Roman", size_pt=12, bold=False)

    # 页眉页脚样式（5号）
    _set_style_font(doc, ["Header", "页眉"], cn="宋体", en="Times New Roman", size_pt=10.5, bold=False)
    _set_style_font(doc, ["Footer", "页脚"], cn="宋体", en="Times New Roman", size_pt=10.5, bold=False)


def main():
    doc = Document()
    _setup_section_page(doc.sections[0])
    configure_styles(doc)

    # 第1节：题名页（不显示页码）
    # 题名页
    add_title(doc, "基于信息蒸馏与注意力融合的轻量图像超分辨率方法研究", 18)
    add_paragraph(doc, "")
    add_paragraph(
        doc,
        "（封面、声明、授权页按学校统一模板另行填写）",
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        first_line_indent_chars=0
    )

    # 第2节：前置部分（从中文摘要开始，罗马页码）
    front_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _setup_section_page(front_sec)
    _set_section_page_number_format(front_sec, fmt="upperRoman", start=1)
    _add_page_number_to_footer(doc, front_sec)

    # 中文摘要
    add_title(doc, "摘  要", 18)
    add_paragraph(doc, "【研究目的】在轻量级单图像超分辨率任务中，如何在有限计算预算下提高重建质量。")
    add_paragraph(doc, "【研究方法】本文提出PFDN网络，包含PFDB信息提炼模块、ESA空间注意力机制以及三路特征融合重加权策略。")
    add_paragraph(doc, "【实验结果】在DIV2K训练、Set5/Set14/B100/Urban100/Manga109测试条件下，与主流轻量模型相比取得更优或可比的精度-复杂度折中。")
    add_paragraph(doc, "【研究结论】所提方法在参数量与推理效率受限场景具有应用价值。")
    add_keywords(doc, cn=True)

    add_chapter_page_break(doc)

    # 英文摘要
    add_title(doc, "Abstract", 18)
    add_paragraph(doc, "This thesis focuses on lightweight single image super-resolution under constrained computation.")
    add_paragraph(doc, "A PFDN model is proposed with PFDB, ESA-based spatial attention, and tri-branch feature fusion with reweighting.")
    add_paragraph(doc, "Experiments on standard benchmarks demonstrate a favorable trade-off between reconstruction quality and efficiency.")
    add_keywords(doc, cn=False)

    add_chapter_page_break(doc)

    # 符号表
    add_title(doc, "符号表（缩略语）", 18)
    add_paragraph(doc, "SISR    单图像超分辨率", first_line_indent_chars=0)
    add_paragraph(doc, "PSNR    峰值信噪比", first_line_indent_chars=0)
    add_paragraph(doc, "SSIM    结构相似性", first_line_indent_chars=0)
    add_paragraph(doc, "FLOPs   浮点运算量", first_line_indent_chars=0)
    add_paragraph(doc, "PFDB    渐进式特征蒸馏块", first_line_indent_chars=0)
    add_paragraph(doc, "ESA     增强型空间注意力", first_line_indent_chars=0)

    add_chapter_page_break(doc)

    # 目录（到二级）
    add_title(doc, "目  录", 18)
    _insert_toc_field(doc, max_level=2)

    # 正文：新节，阿拉伯页码从1开始
    main_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _setup_section_page(main_sec)
    _set_section_page_number_format(main_sec, fmt="decimal", start=1)
    _add_page_number_to_footer(doc, main_sec)

    # 第1章
    add_heading_lv1(doc, "第1章 绪论")
    add_heading_lv2(doc, "1.1 研究背景与意义")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv2(doc, "1.2 国内外研究现状")
    add_heading_lv3(doc, "1.2.1 基于卷积的超分辨率方法")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv3(doc, "1.2.2 轻量化与信息蒸馏方法")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv2(doc, "1.3 研究内容与技术路线")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv2(doc, "1.4 本文主要贡献")
    add_paragraph(doc, "（1）【此处填写】", first_line_indent_chars=0)
    add_paragraph(doc, "（2）【此处填写】", first_line_indent_chars=0)
    add_paragraph(doc, "（3）【此处填写】", first_line_indent_chars=0)
    add_heading_lv2(doc, "1.5 论文结构安排")
    add_paragraph(doc, "【此处填写】")

    add_chapter_page_break(doc)
    add_heading_lv1(doc, "第2章 相关理论与关键技术")
    add_heading_lv2(doc, "2.1 单图像超分辨率退化模型")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv2(doc, "2.2 轻量化网络设计方法")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv2(doc, "2.3 注意力机制基础")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv2(doc, "2.4 评价指标")
    add_paragraph(doc, "【此处填写】")

    add_chapter_page_break(doc)
    add_heading_lv1(doc, "第3章 PFDN轻量级超分网络设计")
    add_heading_lv2(doc, "3.1 网络总体架构")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv2(doc, "3.2 PFDB模块设计")
    add_heading_lv3(doc, "3.2.1 通道分支卷积")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv3(doc, "3.2.2 通道交互融合与局部残差")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv2(doc, "3.3 ESA空间注意力机制")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv2(doc, "3.4 三路特征融合与重加权")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv2(doc, "3.5 上采样重建模块")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv2(doc, "3.6 模型复杂度分析")
    add_paragraph(doc, "【此处填写】")

    add_chapter_page_break(doc)
    add_heading_lv1(doc, "第4章 实验与结果分析")
    add_heading_lv2(doc, "4.1 实验环境与实现细节")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv2(doc, "4.2 数据集与评测协议")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv2(doc, "4.3 对比实验")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv2(doc, "4.4 消融实验")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv2(doc, "4.5 可视化与主观质量分析")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv2(doc, "4.6 复杂度与速度分析")
    add_paragraph(doc, "【此处填写】")

    add_chapter_page_break(doc)
    add_heading_lv1(doc, "第5章 结论与展望")
    add_heading_lv2(doc, "5.1 研究结论")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv2(doc, "5.2 创新点总结")
    add_paragraph(doc, "【此处填写】")
    add_heading_lv2(doc, "5.3 局限性与未来工作")
    add_paragraph(doc, "【此处填写】")

    add_chapter_page_break(doc)
    add_heading_lv1(doc, "参考文献")
    add_paragraph(doc, "[1] 作者. 题名. 期刊名, 年份, 卷(期): 页码.", first_line_indent_chars=0)

    add_chapter_page_break(doc)
    add_heading_lv1(doc, "附录")
    add_paragraph(doc, "附录1 关键代码片段", first_line_indent_chars=0)

    add_chapter_page_break(doc)
    add_heading_lv1(doc, "后记")
    add_paragraph(doc, "【此处填写】")

    _enable_update_fields_on_open(doc)

    doc.save(OUTPUT)
    print(f"已生成：{OUTPUT}")


if __name__ == "__main__":
    main()